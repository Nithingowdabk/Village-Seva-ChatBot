import requests
from bs4 import BeautifulSoup
import openpyxl
import os
from docx import Document

def scrape_toi_topic(topic):
    url = f"https://timesofindia.indiatimes.com/topic/{topic}"
    headers = {"User-Agent": "Mozilla/5.0"}
    
    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        print("Failed to retrieve the page")
        return
    
    # Parse the HTML
    soup = BeautifulSoup(response.text, "html.parser")
    
    # Extract blog titles and URLs
    articles = soup.find_all("span")  # Adjust the tag if needed
    titles = []
    
    for article in articles:
        title = article.text.strip()
        parent = article.find_parent("a")
        if title and parent and parent.has_attr("href"):
            url = "" + parent["href"]
            titles.append((title, url))
    
    # Save titles and URLs in an Excel sheet
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Blog Titles"
    sheet.append(["Title", "URL"])
    
    for title, url in titles:
        sheet.append([title, url])
    
    excel_filename = "toi_blogs.xlsx"
    workbook.save(excel_filename)
    print("Scraping completed. Data saved in", excel_filename)
    
    # Scrape blog articles and save them in doc format
    scrape_blog_articles(excel_filename)

def scrape_blog_articles(excel_file):
    workbook = openpyxl.load_workbook(excel_file)
    sheet = workbook.active
    
    # Create directory for storing blog data
    folder_name = "blogsdata"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    
    for row in sheet.iter_rows(min_row=7, values_only=True):  # Ignoring first 5 URLs
        title, url = row
        if not url:
            continue
        
        response = requests.get(url, headers={"User-Agent": "Mozilla/5.0"})
        if response.status_code != 200:
            print(f"Failed to retrieve {url}")
            continue
        
        soup = BeautifulSoup(response.text, "html.parser")
        article_content = soup.find("div", class_="_s30J clearfix")  # Updated class name
        text = article_content.get_text(strip=True) if article_content else "Content not found"
        
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)
        
        filename = os.path.join(folder_name, f"{title[:50].replace(' ', '_')}.docx")  # Limit filename length
        doc.save(filename)
        print(f"Saved: {filename}")

if __name__ == "__main__":
    user_topic = input("Enter the topic: ")
    scrape_toi_topic(user_topic) 